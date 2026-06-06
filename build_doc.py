"""Generate the Job Application Tracker technical reference (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATE = "June 4, 2026"
PROJECT = "Job Application Tracker"

doc = Document()

# ---- Page setup: US Letter, 1 inch margins ----
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))

# ---- Base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def shade(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code(text):
    """Add a Courier New code block with light shading."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.left_indent = Inches(0.1)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        # ensure the monospace font sticks for all script ranges
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), "Courier New")
        if i < len(lines) - 1:
            run.add_break()
    shade(p)
    return p


def mono(paragraph, text):
    """Append an inline Courier New run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    return run


def body(text):
    return doc.add_paragraph(text)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            mono_cell = isinstance(val, tuple) and val[1] == "mono"
            text = val[0] if isinstance(val, tuple) else val
            run = p.add_run(text)
            run.font.size = Pt(9)
            if mono_cell:
                run.font.name = "Courier New"
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ===================== COVER PAGE =====================
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(PROJECT)
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Personal Technical Reference")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Python CLI · Google OAuth · Gmail API · Google Sheets API · Gemini")
r.font.size = Pt(11)
r.italic = True

for _ in range(8):
    doc.add_paragraph()
d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run(DATE)
r.font.size = Pt(13)

doc.add_page_break()

# ===================== 1. PROJECT OVERVIEW =====================
h1("1. Project Overview")
body(
    "Job Application Tracker is a terminal-only (CLI) assistant that manages a job "
    "search end to end. It stores every application in a Google Sheet, reads Gmail "
    "to automatically discover new applications and advance their status (Applied → "
    "In Review → Interview Scheduled → Rejected/Offer), drafts replies to recruiters "
    "for human approval before sending, and reports weekly metrics. An optional "
    "macOS launchd job runs the whole update once a day."
)
body(
    "The design keeps each module single-purpose and routes all agent logic through "
    "three seams — a storage interface (Tracker), a Gmail wrapper, and an LLM wrapper "
    "— so the Google Sheets backend could be swapped (e.g. for SQLite) without "
    "touching the agent. Application status is modeled as a one-way funnel that only "
    "ever moves forward."
)
h2("Tech Stack")
table(
    ["Concern", "Technology", "Notes"],
    [
        ["Language", "Python 3.11+", "Standard library + a few packages"],
        ["CLI", "argparse", "Sub-command parser in cli.py"],
        ["Auth", "google-auth, google-auth-oauthlib", "InstalledAppFlow (Desktop OAuth)"],
        ["Google APIs", "google-api-python-client", "Gmail v1 + Sheets v4"],
        ["LLM", "google-generativeai", "Gemini 2.5 Flash (free tier)"],
        ["Token security", "cryptography (Fernet)", "Encrypts the OAuth token at rest"],
        ["Config / secrets", "python-dotenv", "Reads .env"],
        ["Scheduling", "macOS launchd", "Daily background run"],
        ["Storage", "Google Sheets", "Behind the Tracker abstraction"],
        ["Tests", "pytest", "Fakes + mocks, no network/quota"],
    ],
    widths=[1.6, 2.6, 2.2],
)

# ===================== 2. ENTRY POINTS =====================
h1("2. Entry Points")
body(
    "There is a single entry point: main.py at the project root. It does nothing but "
    "delegate to the CLI package, keeping the launchable script trivial."
)
code('''# main.py
from jobagent.cli import main

if __name__ == "__main__":
    main()''')

h2("What happens on startup")
body("1. The shell runs ")
code('''python main.py <command> [options]      # e.g.  python main.py discover --days 7''')

h2("Command registration")
body(
    "main() builds an argparse parser with one sub-parser per command. Each "
    "sub-command is wired to a handler via set_defaults(func=...), so dispatch is a "
    "single call — no if/elif ladder."
)
code('''def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(prog="jobagent", description="AI Job Application Tracker")
    sub = parser.add_subparsers(dest="command", required=True)

    sub.add_parser("login", ...).set_defaults(func=cmd_login)
    d = sub.add_parser("discover", ...)
    d.add_argument("--days", type=int, default=7, ...)
    d.set_defaults(func=cmd_discover)
    ...                                  # add, list, sync, update, drafts, summary, task, schedule
    return parser

def main():
    parser = build_parser()
    args = parser.parse_args()
    args.func(args)                      # dispatch to the chosen handler''')
body(
    "Registered commands: setup, login, add, list, sync, discover, update, drafts, "
    "summary, task, schedule."
)

h2("How authentication is initialized")
body(
    "Authentication is lazy — it is only triggered by commands that need Google. A "
    "command calls auth.sheets_service() or auth.gmail_service(); both call "
    "get_credentials(), which loads the Fernet-encrypted token from ~/.jobagent/, "
    "refreshes it silently if expired, and builds the API client. If no token "
    "exists, it raises and tells the user to run login."
)
code('''def get_credentials() -> Credentials:
    payload = _load_encrypted()                       # decrypt ~/.jobagent/token.enc
    if payload is None:
        raise RuntimeError("Not logged in. Run: python main.py login")
    creds = Credentials.from_authorized_user_info(payload, config.SCOPES)
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())                      # silent refresh
        _save_encrypted(json.loads(creds.to_json()))
    return creds

def gmail_service():
    return build("gmail", "v1", credentials=get_credentials())''')

h2("How the agent is orchestrated from one command")
body(
    "A command handler is thin: it builds the three seams (a Tracker, a Gmail "
    "service, and an LLM function) and hands them to one orchestration function in "
    "agent.py. For example, cmd_update wires everything and calls run_discover, "
    "which calls agent.discover_applications. The agent never imports Google or "
    "Gemini directly — it receives them, which is what makes it testable."
)
code('''def cmd_update(args):
    if not _require_gemini():
        return
    run_discover(auth.gmail_service(), _tracker(), args.days, args.max, args.max_llm)

def run_discover(gmail, tracker, days, max_fetch, max_llm):
    refs, emails = _fetch(gmail, days, max_fetch, query=config.APPLICATION_QUERY)
    seen = agent.load_seen()
    summary = agent.discover_applications(
        tracker, emails, llm.extract_application,        # <- injected dependencies
        max_llm=max_llm, seen=seen, apply_keyword_filter=False)
    agent.save_seen(seen)''')

# ===================== 3. FILE DEPENDENCY MAP =====================
h1("3. File Dependency Map")
body(
    "Internal dependencies only (standard-library and third-party imports omitted "
    "for clarity). Note that agent.py depends on the Tracker abstraction by duck "
    "typing — it receives tracker objects rather than importing tracker.py."
)
table(
    ["File", "Imports (internal)", "Depended on by", "Role (one line)"],
    [
        [("main.py", "mono"), ("jobagent.cli", "mono"), "—", "Executable entry point; calls cli.main()"],
        [("jobagent/__init__.py", "mono"), "—", "(package)", "Marks jobagent as a package"],
        [("config.py", "mono"), "—", "auth, tracker, agent, llm, scheduler, setup, cli",
         "Env loading, OAuth scopes, statuses, paths, Gmail queries"],
        [("auth.py", "mono"), ("config", "mono"), "cli, setup",
         "Google OAuth + Fernet-encrypted token store; builds API clients"],
        [("tracker.py", "mono"), ("config", "mono"), "cli (and agent via duck typing)",
         "Application dataclass, Tracker ABC, SheetsTracker"],
        [("gmail_client.py", "mono"), "—", "cli",
         "Thin Gmail wrapper: list / get / parse / send"],
        [("llm.py", "mono"), ("config", "mono"), "cli",
         "Gemini wrapper: classify, extract, propose reply, daily task"],
        [("agent.py", "mono"), ("config", "mono"), "cli",
         "Orchestration: discover, sync, drafts, weekly summary; seen-tracking"],
        [("dates.py", "mono"), "—", "cli", "Parse '30d'/'2w' intervals into a cutoff datetime"],
        [("setup.py", "mono"), ("auth, config", "mono"), "cli", "Guided browser-based first-run wizard"],
        [("scheduler.py", "mono"), ("config", "mono"), "cli", "Install/remove the macOS launchd daily job"],
        [("cli.py", "mono"), ("agent, auth, config, gmail_client, llm, scheduler, setup, dates, tracker", "mono"),
         "main.py", "Registers commands and orchestrates each one"],
    ],
    widths=[1.25, 1.9, 1.7, 1.9],
)

# ===================== 4. DATA FLOW =====================
h1("4. Data Flow")
body(
    "Each major command is traced from CLI input to the external API and back, with "
    "the real code at each step and the data shape entering/leaving it."
)

# ---- 4.1 add ----
h2("4.1  add — manually record an application")
h3("Step 1 — CLI handler (cli.py)")
body("Input: parsed argparse Namespace.")
code('''Namespace(command="add", company="Acme", role="QA Engineer",
          source="LinkedIn", notes=None)''')
code('''def cmd_add(args):
    app = _tracker().add_application(
        args.company, args.role, args.source, notes=args.notes or "")
    print(f"Added: {app.company} — {app.role} [{app.status}]")''')
body("_tracker() builds a SheetsTracker, creating the sheet on first run:")
code('''def _tracker() -> SheetsTracker:
    service = auth.sheets_service()
    sid = get_or_create_sheet(service, config.SPREADSHEET_ID)
    return SheetsTracker(service, sid)''')

h3("Step 2 — SheetsTracker.add_application (tracker.py)")
body("Builds an Application, then appends a row via the Sheets API.")
code('''def add_application(self, company, role, source, date_applied=None, notes=""):
    app = Application(company=company, role=role,
        date_applied=date_applied or _now_iso()[:10],
        source=source, status="Applied",
        last_updated=_now_iso(), notes=notes)
    body = {"values": [[app.company, app.role, app.date_applied, app.source,
                        app.status, app.last_updated, app.notes]]}
    self.service.spreadsheets().values().append(
        spreadsheetId=self.spreadsheet_id, range=self.tab,
        valueInputOption="USER_ENTERED", body=body).execute(num_retries=3)
    return app''')
body("Data leaving Python → Sheets API (one row, column order = SHEET_HEADERS):")
code('''{"values": [["Acme", "QA Engineer", "2026-06-04", "LinkedIn",
             "Applied", "2026-06-04T09:00:00+00:00", ""]]}''')
body("Returned object:")
code('''Application(company="Acme", role="QA Engineer", date_applied="2026-06-04",
            source="LinkedIn", status="Applied",
            last_updated="2026-06-04T09:00:00+00:00", notes="")''')

# ---- 4.2 list ----
h2("4.2  list — show applications with filters")
h3("Step 1 — CLI handler (cli.py)")
code('''def cmd_list(args):
    since = parse_since(args.since)                 # "30d" -> datetime cutoff or None
    apps = _tracker().get_applications(since=since, status=args.status)''')
h3("Step 2 — parse_since (dates.py)")
code('''def parse_since(text, now=None):
    if not text: return None
    now = now or datetime.now(timezone.utc)
    m = re.fullmatch(r"(\\d+)\\s*([dw])", text.strip().lower())
    if not m: raise ValueError(...)
    return now - timedelta(days=int(m.group(1)) * {"d":1,"w":7}[m.group(2)])''')
body('Transformation: ')
p = doc.paragraphs[-1]
mono(p, '"30d"'); p.add_run(" → "); mono(p, "datetime(2026-05-05, tzinfo=UTC)")
h3("Step 3 — SheetsTracker.get_applications (tracker.py)")
body("Reads all rows, maps each to an Application, filters in Python.")
code('''def get_applications(self, since=None, status=None):
    rows = self._all_rows()                         # Sheets values().get()
    apps = [self._to_app(r) for r in rows[1:]]      # skip header row
    if status: apps = [a for a in apps if a.status == status]
    if since:  apps = [a for a in apps if a.date_applied >= since.date().isoformat()]
    return apps''')
body("Data shape from Sheets API → list of raw rows → list[Application]:")
code('''# in:  {"values": [["Company","Role",...], ["Acme","QA Engineer",...]]}
# out: [Application(company="Acme", role="QA Engineer", status="Applied", ...)]''')

# ---- 4.3 discover / update ----
h2("4.3  discover / update — auto-populate & advance from Gmail")
body(
    "The flagship flow. update is discover run on a wide window; both call the same "
    "code. Trace: CLI → Gmail search → fetch+parse → Gemini extract → per-company "
    "merge → Sheets upsert."
)
h3("Step 1 — CLI handler + fetch (cli.py)")
code('''def run_discover(gmail, tracker, days, max_fetch, max_llm):
    refs, emails = _fetch(gmail, days, max_fetch, query=config.APPLICATION_QUERY)
    seen = agent.load_seen()
    summary = agent.discover_applications(
        tracker, emails, llm.extract_application,
        max_llm=max_llm, seen=seen, apply_keyword_filter=False)
    agent.save_seen(seen)

def _fetch(gmail, days, max_results, query=None):
    refs = gmail_client.list_recent(gmail, days=days, max_results=max_results, query=query)
    return refs, [gmail_client.parse_message(gmail_client.get_message(gmail, r["id"]))
                  for r in refs]''')
h3("Step 2 — Gmail search (gmail_client.list_recent)")
body("Builds a Gmail query string and lists matching message ids.")
code('''def list_recent(service, days=2, max_results=100, query=None):
    q = f"newer_than:{days}d"
    if query: q += f" ({query})"
    resp = service.users().messages().list(
        userId="me", q=q, maxResults=max_results).execute(num_retries=3)
    return resp.get("messages", [])''')
body("Query sent to Gmail (date window AND the precise ATS query):")
code('''newer_than:7d (from:greenhouse.io OR ... OR subject:"thank you for applying" OR ...)''')
body("Gmail returns message references:")
code('''[{"id": "18f...a1", "threadId": "18f...a1"}, {"id": "18f...b2", ...}]''')
h3("Step 3 — Fetch + flatten each message (gmail_client.parse_message)")
code('''def parse_message(msg):
    payload = msg.get("payload", {})
    headers = payload.get("headers", [])
    sender = _header(headers, "From")
    return {"id": msg.get("id",""), "thread_id": msg.get("threadId",""),
            "sender": sender, "sender_domain": _domain(sender),
            "subject": _header(headers,"Subject"),
            "snippet": msg.get("snippet",""), "body": _extract_plain(payload)}''')
body("Raw Gmail message (nested) → flat dict the agent understands:")
code('''# in (raw):  {"id":"18f..","payload":{"headers":[{"name":"From","value":"GitLab <no-reply@greenhouse.io>"},
#                                        {"name":"Subject","value":"Thank you for applying to GitLab"}],
#                          "parts":[{"mimeType":"text/plain","body":{"data":"<base64>"}}]}, "snippet":"..."}
# out (flat): {"id":"18f..","thread_id":"18f..","sender":"GitLab <no-reply@greenhouse.io>",
#              "sender_domain":"greenhouse.io","subject":"Thank you for applying to GitLab",
#              "snippet":"...","body":"Hi Hiya, thanks for applying ..."}''')
h3("Step 4 — Gemini extraction (llm.extract_application)")
body("One Gemini call per email; parses strict JSON; validates status.")
code('''def extract_application(email):
    prompt = ("You read a single email and decide if it concerns a job the RECIPIENT "
              "applied to ... Respond with ONLY a JSON object: is_job_application (bool), "
              "company, role, status (one of: Applied, In Review, Interview Scheduled, "
              "Rejected, Offer), confidence (0.0-1.0).")
    try:
        data = _extract_json(_generate(prompt))     # Gemini call + JSON parse
        status = data.get("status")
        if status not in config.STATUSES: status = "Applied"
        return {"is_job_application": bool(data.get("is_job_application", False)),
                "company": data.get("company"), "role": data.get("role"),
                "status": status, "confidence": float(data.get("confidence", 0.0))}
    except (ValueError, json.JSONDecodeError, TypeError):
        return {"is_job_application": False, "company": None, "role": None,
                "status": "Applied", "confidence": 0.0}''')
body("flat email dict → Gemini → structured extraction:")
code('''# out: {"is_job_application": true, "company": "GitLab",
#        "role": "Backend Engineer", "status": "Applied", "confidence": 0.93}''')
h3("Step 5 — Per-company merge + Sheets upsert (agent.discover_applications)")
body(
    "Skips noise/seen emails, caps LLM calls, then collapses every email from one "
    "company into a single application (best role + most-advanced status), and "
    "adds or forward-advances the Sheet row."
)
code('''candidates = [e for e in emails
              if (not apply_keyword_filter or _is_job_candidate(e))
              and not _is_noise(e) and e.get("id") not in seen]
to_scan = candidates[:max_llm]
merged = {}
for email in to_scan:
    r = extractor(email)                             # llm.extract_application
    seen.add(email.get("id"))
    if not r["is_job_application"] or r["confidence"] < min_confidence: continue
    key = r["company"].lower()
    ...  # keep a real role over "(unknown)", keep most-advanced status via _advances()

for m in merged.values():
    existing = tracker.find_by_company(m["company"])  # one row per company
    if existing is None:
        tracker.add_application(m["company"], m["role"], source="Email")
        if _advances(m["status"], "Applied"):
            tracker.update_status(m["company"], m["role"], m["status"], note="...")
    elif _advances(m["status"], existing.status):
        tracker.update_status(existing.company, existing.role, m["status"], note="...")
return {"added": [...], "updated": [...], "skipped_quota": N, "error": None}''')
body("Three GitLab emails (OTP, confirmation, rejection) merge to one result:")
code('''# merged["gitlab"] = {"company":"GitLab","role":"Backend Engineer","status":"Rejected"}
# -> Sheet row GitLab advances Applied -> Rejected (status only moves forward)''')

# ---- 4.4 sync ----
h2("4.4  sync — update statuses of already-tracked companies")
h3("Step 1 — fetch recent mail, classify each (cli.run_sync → agent.sync_inbox)")
code('''def sync_inbox(tracker, emails, classifier, min_confidence=0.6, log_fn=_log_change):
    apps = tracker.get_applications()
    candidates = [(a.company, a.role) for a in apps]
    for email in _prefilter(emails, apps):           # cheap: only emails naming a tracked co.
        result = classifier(email, candidates)        # llm.classify_email (Gemini)
        new_status, company = result.get("new_status"), result.get("matched_company")
        if not (new_status and company): continue
        if result["confidence"] < min_confidence: continue
        existing = tracker.find_by_company(company)
        if existing is None or not _advances(new_status, existing.status): continue
        tracker.update_status(existing.company, existing.role, new_status, note="...")''')
h3("Step 2 — classify_email maps intent → status (llm.py)")
code('''INTENT_STATUS = {"confirmation":"In Review", "rejection":"Rejected",
                 "interview_invite":"Interview Scheduled", "offer":"Offer",
                 "recruiter_followup":None, "other":None}
# Gemini returns intent; Python (not the model) decides the status:
# {"matched_company":"Acme","matched_role":"QA Engineer",
#  "intent":"interview_invite","new_status":"Interview Scheduled","confidence":0.9}''')

# ---- 4.5 drafts ----
h2("4.5  drafts — AI reply drafting with human approval")
h3("Step 1 — fetch reply-worthy mail + propose (cli.cmd_drafts)")
code('''refs, emails = _fetch(gmail, args.days, args.max, query=config.REPLY_QUERY)
seen = agent.load_seen(config.DRAFTS_SEEN_FILE)
result = agent.propose_drafts(emails, llm.propose_reply, seen=seen, max_llm=args.max_llm)''')
h3("Step 2 — one call decides need-reply AND drafts it (llm.propose_reply)")
code('''def propose_reply(email):
    prompt = ("...Decide if it needs a PERSONAL reply ... If it does, draft a concise, "
              "warm, professional reply under 150 words. Respond with ONLY JSON: "
              '{"needs_reply": bool, "draft": string}.')
    data = _extract_json(_generate(prompt))
    return {"needs_reply": bool(data.get("needs_reply", False)),
            "draft": (data.get("draft") or "").strip()}
# out: {"needs_reply": true, "draft": "Hi Sarah, thank you for reaching out ..."}''')
h3("Step 3 — human approves, then send (agent.review_draft → gmail_client.send_message)")
code('''def review_draft(draft, prompt_fn, edit_fn):
    choice = prompt_fn("[a]pprove & send / [e]dit / [s]kip? ").strip().lower()
    if choice == "a": return draft
    if choice == "e": return edit_fn(draft)         # opens $EDITOR
    return None                                     # skip -> nothing sent

# only on approval:
to = gmail_client.sender_email(email["sender"])     # "Sarah <s@acme.com>" -> "s@acme.com"
gmail_client.send_message(gmail, to, f"Re: {email['subject']}", final,
                          thread_id=email.get("thread_id"))''')
body("send_message builds a MIME message, base64url-encodes it, and posts to Gmail:")
code('''# {"raw": "<base64url MIME>", "threadId": "18f...a1"}  ->  users().messages().send()''')

# ---- 4.6 summary ----
h2("4.6  summary — weekly metrics (no LLM)")
code('''def weekly_summary(tracker, now=None):
    now = now or datetime.now(timezone.utc)
    week_ago = (now - timedelta(days=7)).date().isoformat()
    apps = tracker.get_applications()
    responded = [a for a in apps if a.status != "Applied"]
    by_status = {s: sum(1 for a in apps if a.status == s) for s in config.STATUSES}
    return {"total": len(apps), "by_status": by_status, "responses": len(responded),
            "response_rate": round(len(responded)/len(apps), 2) if apps else 0.0,
            "interviews": by_status["Interview Scheduled"], "offers": by_status["Offer"],
            "rejections": by_status["Rejected"],
            "pending": by_status["Applied"] + by_status["In Review"],
            "applied_this_week": sum(1 for a in apps if a.date_applied >= week_ago)}''')
body("list[Application] → metrics dict (pure computation, read straight from the Sheet):")
code('''# {"total":4,"responses":2,"response_rate":0.5,"interviews":0,"offers":0,
#  "rejections":1,"pending":3,"applied_this_week":4, "by_status":{...}}''')

# ===================== 5. ARCHITECTURE DIAGRAM =====================
h1("5. Architecture Diagram")
body("Arrows are labelled with the data flowing across them.")
code('''                                  +-------------------+
                                  |     main.py       |   executable entry point
                                  +---------+---------+
                                            | args.func(args)
                                            v
+---------------------------------------------------------------------------------+
|                              CLI LAYER  (cli.py)                                 |
|   argparse sub-commands -> cmd_* handlers; builds the 3 seams, prints output    |
+----+--------------------------+----------------------------+--------------------+
     | Tracker                  | gmail service              | llm functions
     | (SheetsTracker)          | (Gmail client)             | (Gemini)
     v                          v                            v
+----------------+   +------------------------+   +--------------------------------+
| ORCHESTRATION  |   |     auth.py            |   |          llm.py                |
|  agent.py      |   |  OAuth + Fernet token  |   | classify / extract / reply /   |
| discover/sync/ |   |  get_credentials()     |   | daily_task   (JSON in/out)     |
| drafts/summary |   +-----------+------------+   +---------------+----------------+
+--+----+----+---+               |  credentials                 | prompt / JSON
   |    |    |                   |                               |
   |    |    | Application objs  |                               |
   |    |    v                   |                               |
   |    | +-----------------+    |                               |
   |    | |  tracker.py     |    |                               |
   |    | | Tracker (ABC)   |    |                               |
   |    | | SheetsTracker   |    |                               |
   |    | +--------+--------+    |                               |
   |    |          | rows/values |                               |
   | flat email    |             |                               |
   | dicts         |             |                               |
   v               v             v                               v
+--------------+  +-----------------------+   +----------------------+   +-----------------+
| gmail_client |  |   Google Sheets API   |   |     Gmail API        |   |   Gemini API    |
| list/get/    |  |  (spreadsheets.values |   | (users.messages      |   | (generateContent)|
| parse/send   |->|   get/append/update)  |   |  list/get/send)      |   |                 |
+------+-------+  +-----------------------+   +----------+-----------+   +--------+--------+
       |  message refs / MIME send                      ^                         ^
       +------------------------------------------------+                         |
                       reads & sends mail                                         |
   config.py  ......... feeds scopes, queries, paths, model name to every layer .. (cross-cutting)
   dates.py / setup.py / scheduler.py ... support: interval parsing, wizard, launchd daily run

  External services:   [ Gmail API ]   [ Google Sheets API ]   [ Gemini API ]
  Data on the wire:    emails (JSON)   rows (values arrays)     prompts/JSON''')

# ===================== 6. KEY DESIGN DECISIONS =====================
h1("6. Key Design Decisions")

h2("Storage behind a Tracker abstraction")
body(
    "tracker.py defines an abstract Tracker plus an Application dataclass; "
    "SheetsTracker is the only concrete implementation today. Agent logic depends "
    "on this interface (by duck typing — it receives a tracker), never on the Sheets "
    "API. Swapping to SQLite means writing one new class and changing one wiring "
    "line in cli._tracker(). The rest of the code only ever sees Application objects, "
    "not raw spreadsheet rows."
)

h2("Dependency injection at the seams")
body(
    "agent.py functions accept the tracker, the email list, and the LLM function as "
    "parameters (e.g. discover_applications(tracker, emails, extractor, ...)). This "
    "is why the test suite can exercise the full logic with an in-memory FakeTracker "
    "and a stub extractor — no network, no quota. The CLI is the only place that "
    "wires real Google/Gemini objects in."
)

h2("Status as a one-way funnel")
body(
    "Statuses are ranked (Applied < In Review < Interview Scheduled < Rejected < "
    "Offer) and a helper _advances() guarantees status only ever moves forward. "
    "Both discover and sync use it, so a stray or out-of-order email can never "
    "downgrade a status (e.g. knock a Rejected role back to In Review)."
)

h2("One row per company")
body(
    "A company sends many emails (OTP, confirmation, interview, rejection). The "
    "agent merges them by company — keeping a real role over '(unknown)' and the "
    "most-advanced status — and the tracker matches with find_by_company(). This "
    "collapses the funnel into a single, correct row instead of duplicates. Trade-"
    "off: two genuinely different roles at one company would merge into one row."
)

h2("Aggressive quota economy (the LLM is the scarce resource)")
body(
    "The Gemini free tier allows only ~5 requests/minute and a small daily cap, so "
    "the design spends LLM calls carefully: (1) Gmail's own search pre-filters to "
    "real application mail for free; (2) OTP/verification emails are dropped before "
    "any call; (3) scanned email ids are persisted in seen.json so an email is read "
    "by the model at most once, ever; (4) discover/extract is a single pass that "
    "both adds and advances, avoiding a second classification pass; (5) propose_reply "
    "decides need-reply and drafts in one call. _call_with_retry() backs off on the "
    "per-minute limit but fails fast when the daily cap is gone."
)

h2("Encrypted credentials & least-privilege scopes")
body(
    "The OAuth token is stored encrypted with Fernet; the key lives in a separate "
    "0600 file under ~/.jobagent/. Only the scopes the app actually uses are "
    "requested (gmail.readonly, gmail.send, spreadsheets, profile). Secrets (.env, "
    "client_secret.json, tokens) are gitignored and never committed."
)

h2("Human-in-the-loop for anything outbound")
body(
    "Discovery and status updates are automatic, but the agent never sends an email "
    "on its own. review_draft() requires an explicit approve before send_message() "
    "is called, and the daily scheduled job runs only update (read + record), never "
    "drafts."
)

h2("Resilience to flaky networks and bad model output")
body(
    "Every Google API call passes num_retries=3 so transient timeouts/5xx are "
    "retried with backoff. LLM responses are parsed defensively with _extract_json "
    "(tolerates code fences) and every llm.* function returns a safe default on "
    "parse failure, so one malformed response never crashes a run."
)

doc.save("Job_Application_Tracker_Technical_Reference.docx")
print("Saved Job_Application_Tracker_Technical_Reference.docx")
